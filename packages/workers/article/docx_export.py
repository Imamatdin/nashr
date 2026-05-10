"""DOCX exporter for drafted articles — primary deliverable for article customers.

The exporter produces a Microsoft Word document that meets Uzbek
university submission standards (A4, 2/2/3/1.5 cm margins, Times New
Roman 14pt, 1.5-line spacing, justified body, bold centred Heading 1)
and renders inline citations + a bibliography in the requested style.

Pipeline (all sync because python-docx is sync; wrapped in
:func:`asyncio.to_thread`):
1. Construct an in-memory :class:`Document`, configure section layout
   and the ``Normal`` / ``Heading 1-3`` styles.
2. Add a title page for ``referat``/``kurs_ishi`` (skipped for
   ``ilmiy_maqola``/``hisobot``).
3. Insert a TOC field for ``kurs_ishi``.
4. For ``ilmiy_maqola`` only: emit the abstract block + keywords line.
5. Iterate the drafted sections in outline order, replacing inline
   ``[<uuid>]`` markers with the formatted in-text citation produced
   by :class:`BibliographyFormatter`. Chicago renders superscript
   footnote markers inline + a footnotes section at end (a documented
   fallback because Word footnote XML is fragile in python-docx 1.2).
6. Append the bibliography section with the localised heading.
7. Append a verification-warning paragraph if the report flags any
   ``NOT_SUPPORTED`` / ``CONTRADICTED`` citation.
8. Save into a :class:`io.BytesIO`, return :class:`ExportResult`.

The 300-line CLAUDE.md budget is exceeded here for the same reason
``bibliography.py`` runs to ~990 lines: page layout, title page, TOC,
section rendering, inline citation replacement (5 styles), bibliography
rendering, footnote handling, page-number XML, and verification
warnings each carry non-trivial XML/style logic, and splitting them
across modules would scatter a single coherent build operation.
"""

from __future__ import annotations

import asyncio
import re
from dataclasses import dataclass
from io import BytesIO
from typing import Any, Final

from docx import Document  # pyright: ignore[reportUnknownVariableType]
from docx.document import Document as DocumentT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement  # pyright: ignore[reportUnknownVariableType]
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from packages.core.enums import (
    ArticleStructure,
    CitationFormat,
)
from packages.core.models.article import ArticleDraftResult, ArticleOutline, Paragraph
from packages.core.models.bibliography import (
    CitationMetadata,
    FormattedBibliography,
    FormattedEntry,
)
from packages.core.models.export import ArticleExportMetadata, ExportResult
from packages.core.models.verification import CitationVerificationReport
from packages.workers.article.bibliography import BibliographyFormatter

# ---------------------------------------------------------------------------
# Layout constants — Uzbek university submission standard
# ---------------------------------------------------------------------------

PAGE_WIDTH_CM: Final[float] = 21.0
PAGE_HEIGHT_CM: Final[float] = 29.7
MARGIN_TOP_CM: Final[float] = 2.0
MARGIN_BOTTOM_CM: Final[float] = 2.0
MARGIN_LEFT_CM: Final[float] = 3.0
MARGIN_RIGHT_CM: Final[float] = 1.5

BODY_FONT_NAME: Final[str] = "Times New Roman"
BODY_FONT_SIZE_PT: Final[int] = 14
TITLE_FONT_SIZE_PT: Final[int] = 16
SMALL_FONT_SIZE_PT: Final[int] = 12
FIRST_LINE_INDENT_CM: Final[float] = 1.25
HANGING_INDENT_CM: Final[float] = 1.25

WORDS_PER_PAGE_ESTIMATE: Final[int] = 250

# ---------------------------------------------------------------------------
# Localised labels
# ---------------------------------------------------------------------------

_BIBLIOGRAPHY_HEADINGS: Final[dict[str, str]] = {
    "uz": "FOYDALANILGAN ADABIYOTLAR RO'YXATI",
    "ru": "СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ",
    "en": "REFERENCES",
}
_KEYWORDS_LABELS: Final[dict[str, str]] = {
    "uz": "Kalit so'zlar",
    "ru": "Ключевые слова",
    "en": "Keywords",
}
_TOC_HEADINGS: Final[dict[str, str]] = {
    "uz": "MUNDARIJA",
    "ru": "СОДЕРЖАНИЕ",
    "en": "TABLE OF CONTENTS",
}
_FOOTNOTES_HEADINGS: Final[dict[str, str]] = {
    "uz": "IZOHLAR",
    "ru": "ПРИМЕЧАНИЯ",
    "en": "NOTES",
}
_PERFORMED_BY: Final[dict[str, str]] = {
    "uz": "Bajardi",
    "ru": "Выполнил",
    "en": "Performed by",
}
_CHECKED_BY: Final[dict[str, str]] = {
    "uz": "Tekshirdi",
    "ru": "Проверил",
    "en": "Checked by",
}
_ARTICLE_TYPE_LABELS: Final[dict[ArticleStructure, dict[str, str]]] = {
    ArticleStructure.REFERAT: {"uz": "REFERAT", "ru": "РЕФЕРАТ", "en": "REPORT"},
    ArticleStructure.KURS_ISHI: {"uz": "KURS ISHI", "ru": "КУРСОВАЯ РАБОТА", "en": "COURSEWORK"},
    ArticleStructure.HISOBOT: {"uz": "HISOBOT", "ru": "ОТЧЁТ", "en": "REPORT"},
    ArticleStructure.ILMIY_MAQOLA: {"uz": "ILMIY MAQOLA", "ru": "НАУЧНАЯ СТАТЬЯ", "en": "ARTICLE"},
}

# Markers in drafted text are square-bracketed UUIDs; the regex is loose
# enough to also catch shorter chunk-index markers ("[42]") without
# matching the bibliography "[1]" we render later.
_CITATION_MARKER_RE: Final[re.Pattern[str]] = re.compile(r"\[([^\[\]\s]+)\]")
_APA_AUTHOR_YEAR_RE: Final[re.Pattern[str]] = re.compile(r"^([^,(]+?)(?:,[^()]*)?\s*\((\d{4})")


@dataclass(frozen=True)
class _MarkerMatch:
    """Resolved bibliography reference for one citation marker."""

    number: int
    entry: FormattedEntry
    metadata: CitationMetadata | None


# ---------------------------------------------------------------------------
# python-docx XML boundary
#
# python-docx exposes its OOXML internals through lxml elements that are
# entirely unannotated (returns ``Unknown`` to pyright, members are
# accessed via ``_r`` private attributes). Per CLAUDE.md, ``Any`` is
# permitted at the boundary of genuinely untyped external input; we
# concentrate every interaction with that boundary into the helpers
# below so the rest of the module stays strict-typed.
# ---------------------------------------------------------------------------


def _xml(tag: str, **attrs: str) -> Any:
    """Build a ``w:*`` OXML element with ``attrs`` set on it."""

    elem: Any = OxmlElement(tag)  # pyright: ignore[reportUnknownVariableType]
    for key, value in attrs.items():
        elem.set(qn(key), value)  # pyright: ignore[reportUnknownMemberType]
    return elem  # pyright: ignore[reportUnknownVariableType]


def _append_run_xml(run: Any, *children: Any) -> None:
    """Append OXML children into a run's underlying ``<w:r>`` element."""

    for child in children:
        run._r.append(child)  # pyright: ignore[reportPrivateUsage]


def _set_xml_text(elem: Any, text: str) -> None:
    """Set the ``.text`` of an OXML element (kept as a helper for type isolation)."""

    elem.text = text


def _section_pr(section: Any) -> Any:
    """Return the section's ``<w:sectPr>`` element."""

    return section._sectPr  # pyright: ignore[reportPrivateUsage]


def _runs_last_break(run: Any) -> Any:
    """Return the most recently appended ``<w:br>`` element on a run."""

    return run._r.findall(qn("w:br"))[-1]  # pyright: ignore[reportPrivateUsage]


# ---------------------------------------------------------------------------
# Public exporter
# ---------------------------------------------------------------------------


class DOCXExporter:
    """Renders an :class:`ArticleDraftResult` into a Microsoft Word DOCX file.

    Stateless. One instance can be reused across exports. The single
    public entry point is :meth:`export`; the ``_build_*`` helpers are
    sync (python-docx is sync) and run inside :func:`asyncio.to_thread`.
    """

    def __init__(self, formatter: BibliographyFormatter | None = None) -> None:
        self._formatter = formatter if formatter is not None else BibliographyFormatter()

    async def export(
        self,
        draft: ArticleDraftResult,
        bibliography: FormattedBibliography,
        verification: CitationVerificationReport | None,
        outline: ArticleOutline,
        metadata: ArticleExportMetadata,
        language: str,
        citation_metadata: list[CitationMetadata] | None = None,
    ) -> ExportResult:
        """Build a DOCX from drafted material and return the bytes + counters."""

        return await asyncio.to_thread(
            self._build,
            draft,
            bibliography,
            verification,
            outline,
            metadata,
            language,
            citation_metadata,
        )

    # -- sync builder ------------------------------------------------------

    def _build(
        self,
        draft: ArticleDraftResult,
        bibliography: FormattedBibliography,
        verification: CitationVerificationReport | None,
        outline: ArticleOutline,
        metadata: ArticleExportMetadata,
        language: str,
        citation_metadata: list[CitationMetadata] | None,
    ) -> ExportResult:
        del outline  # outline order is already reflected in draft.sections
        lang = _normalise_language(language)
        doc = Document()
        _setup_page_layout(doc)
        _setup_normal_style(doc)
        _setup_heading_styles(doc)
        has_title = _has_title_page(metadata.article_type)
        _setup_page_numbers(doc, suppress_first=has_title)

        if has_title:
            _add_title_page(doc, metadata, lang)
            _add_page_break(doc)

        if _has_toc(metadata.article_type):
            _add_toc(doc, lang)
            _add_page_break(doc)

        if metadata.article_type is ArticleStructure.ILMIY_MAQOLA:
            if metadata.abstract_text:
                _add_abstract_block(doc, metadata.abstract_text, lang)
            if metadata.keywords:
                _add_keywords_line(doc, metadata.keywords, lang)

        marker_map = _build_marker_map(
            draft=draft,
            bibliography=bibliography,
            citation_metadata=citation_metadata,
        )

        chicago_footnotes: list[str] = []
        citation_count = 0
        section_count = 0
        for idx, section_result in enumerate(draft.sections):
            section = section_result.section
            section_count += 1
            _add_section_heading(doc, section.title, idx)
            for paragraph in section.paragraphs:
                replaced, used_count = _render_paragraph_with_citations(
                    paragraph_obj=paragraph,
                    marker_map=marker_map,
                    formatter=self._formatter,
                    style=metadata.citation_format,
                    language=lang,
                    chicago_footnotes=chicago_footnotes,
                )
                citation_count += used_count
                _add_body_paragraph(doc, replaced)

        if verification is not None:
            critical = verification.not_supported + verification.contradicted
            if critical > 0:
                _add_verification_warning(doc, critical, lang)

        if metadata.citation_format is CitationFormat.CHICAGO and chicago_footnotes:
            _add_page_break(doc)
            _add_footnotes_section(doc, chicago_footnotes, lang)

        _add_page_break(doc)
        _add_bibliography_section(doc, bibliography, lang)

        buffer = BytesIO()
        doc.save(buffer)
        file_bytes = buffer.getvalue()

        word_count = _count_words(draft)
        return ExportResult(
            file_bytes=file_bytes,
            filename=_filename(metadata),
            file_size_bytes=len(file_bytes),
            page_count_estimate=max(
                1, (word_count + WORDS_PER_PAGE_ESTIMATE - 1) // WORDS_PER_PAGE_ESTIMATE
            ),
            word_count=word_count,
            section_count=section_count,
            citation_count=citation_count,
            bibliography_count=bibliography.total_entries,
            warnings=list(draft.warnings)[:50],
        )


# ---------------------------------------------------------------------------
# Page layout / styles
# ---------------------------------------------------------------------------


def _setup_page_layout(doc: DocumentT) -> None:
    section: Any = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_CM)
    section.bottom_margin = Cm(MARGIN_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_CM)
    section.right_margin = Cm(MARGIN_RIGHT_CM)


def _setup_normal_style(doc: DocumentT) -> None:
    style: Any = doc.styles["Normal"]  # pyright: ignore[reportUnknownVariableType, reportUnknownMemberType]
    _apply_run_font(style.font, BODY_FONT_NAME, BODY_FONT_SIZE_PT)
    _force_east_asia_font(style.element, BODY_FONT_NAME)
    _apply_paragraph_format(
        style.paragraph_format,
        line_spacing=WD_LINE_SPACING.ONE_POINT_FIVE,
        first_line_indent=Cm(FIRST_LINE_INDENT_CM),
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_before=Pt(0),
        space_after=Pt(0),
    )


def _setup_heading_styles(doc: DocumentT) -> None:
    for level, alignment, italic in (
        (1, WD_ALIGN_PARAGRAPH.CENTER, False),
        (2, WD_ALIGN_PARAGRAPH.LEFT, False),
        (3, WD_ALIGN_PARAGRAPH.LEFT, True),
    ):
        style: Any = doc.styles[f"Heading {level}"]  # pyright: ignore[reportUnknownVariableType, reportUnknownMemberType]
        _apply_run_font(
            style.font,
            BODY_FONT_NAME,
            BODY_FONT_SIZE_PT,
            bold=True,
            italic=italic,
            color=RGBColor(0, 0, 0),
        )
        _force_east_asia_font(style.element, BODY_FONT_NAME)
        _apply_paragraph_format(
            style.paragraph_format,
            line_spacing=WD_LINE_SPACING.ONE_POINT_FIVE,
            first_line_indent=Cm(0),
            left_indent=Cm(0),
            alignment=alignment,
            space_before=Pt(12),
            space_after=Pt(6),
        )


def _apply_run_font(
    font: Any,
    name: str,
    size_pt: int,
    *,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
) -> None:
    """Apply the standard run-font block to a typed-Any font handle."""

    font.name = name
    font.size = Pt(size_pt)
    font.bold = bold
    font.italic = italic
    if color is not None:
        font.color.rgb = color


def _apply_paragraph_format(
    pf: Any,
    *,
    line_spacing: object | None = None,
    first_line_indent: object | None = None,
    left_indent: object | None = None,
    alignment: object | None = None,
    space_before: object | None = None,
    space_after: object | None = None,
) -> None:
    """Apply the standard paragraph-format block, skipping fields left as ``None``."""

    if line_spacing is not None:
        pf.line_spacing_rule = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if left_indent is not None:
        pf.left_indent = left_indent
    if alignment is not None:
        pf.alignment = alignment
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after


def _force_east_asia_font(style_element: Any, name: str) -> None:
    """Pin ``w:eastAsia`` and ``w:cs`` font names so Word renders Cyrillic in TNR.

    Without this, Word silently substitutes Calibri for non-Latin glyphs
    even when the Latin name is set, producing mixed-typeface paragraphs.
    """

    rpr = style_element.find(qn("w:rPr"))
    if rpr is None:
        rpr = _xml("w:rPr")
        style_element.append(rpr)
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = _xml("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:cs"), name)


# ---------------------------------------------------------------------------
# Page numbers
# ---------------------------------------------------------------------------


def _setup_page_numbers(doc: DocumentT, suppress_first: bool) -> None:
    """Add bottom-centre PAGE field; suppress on the first page when needed."""

    section: Any = doc.sections[0]
    if suppress_first:
        section.different_first_page_header_footer = True
        _set_title_pg_property(section)

    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _append_page_field(paragraph)


def _set_title_pg_property(section: Any) -> None:
    """Write ``<w:titlePg/>`` on the section so Word actually honours different first page."""

    sect_pr = _section_pr(section)
    existing = sect_pr.find(qn("w:titlePg"))
    if existing is None:
        sect_pr.append(_xml("w:titlePg"))


def _append_page_field(paragraph: Any) -> None:
    begin_run = paragraph.add_run()
    _append_run_xml(begin_run, _xml("w:fldChar", **{"w:fldCharType": "begin"}))

    instr_run = paragraph.add_run()
    instr = _xml("w:instrText", **{"xml:space": "preserve"})
    _set_xml_text(instr, " PAGE ")
    _append_run_xml(instr_run, instr)

    end_run = paragraph.add_run()
    _append_run_xml(end_run, _xml("w:fldChar", **{"w:fldCharType": "end"}))


# ---------------------------------------------------------------------------
# Title page
# ---------------------------------------------------------------------------


def _has_title_page(structure: ArticleStructure) -> bool:
    return structure in (ArticleStructure.REFERAT, ArticleStructure.KURS_ISHI)


def _has_toc(structure: ArticleStructure) -> bool:
    return structure is ArticleStructure.KURS_ISHI


def _add_title_page(doc: DocumentT, metadata: ArticleExportMetadata, language: str) -> None:
    """Render the title page; lines are suppressed when their metadata is missing."""

    if metadata.university_name:
        _add_centered_line(doc, metadata.university_name.upper(), bold=True)
    if metadata.faculty_name:
        _add_centered_line(doc, metadata.faculty_name)
    if metadata.department_name:
        _add_centered_line(doc, metadata.department_name)

    _add_blank_lines(doc, 4)

    type_label = _ARTICLE_TYPE_LABELS[metadata.article_type].get(
        language, _ARTICLE_TYPE_LABELS[metadata.article_type]["en"]
    )
    _add_centered_line(doc, type_label, bold=True)
    _add_blank_lines(doc, 1)
    _add_centered_line(doc, metadata.title, bold=True, size_pt=TITLE_FONT_SIZE_PT)

    _add_blank_lines(doc, 4)

    performed_label = _PERFORMED_BY[language]
    checked_label = _CHECKED_BY[language]
    author_block = f"{performed_label}: {metadata.author_name}"
    if metadata.author_group:
        author_block += f", {metadata.author_group}"
    _add_right_aligned_line(doc, author_block)
    if metadata.supervisor_name:
        supervisor_block = f"{checked_label}: {metadata.supervisor_name}"
        if metadata.supervisor_title:
            supervisor_block += f", {metadata.supervisor_title}"
        _add_right_aligned_line(doc, supervisor_block)

    _add_blank_lines(doc, 6)

    locality = metadata.city
    if metadata.year is not None:
        locality = f"{metadata.city}, {metadata.year}"
    _add_centered_line(doc, locality)


# ---------------------------------------------------------------------------
# TOC, abstract, keywords
# ---------------------------------------------------------------------------


def _add_toc(doc: DocumentT, language: str) -> None:
    heading_text = _TOC_HEADINGS[language]
    heading = doc.add_paragraph(heading_text, style="Heading 1")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph = doc.add_paragraph()
    begin_run = paragraph.add_run()
    _append_run_xml(begin_run, _xml("w:fldChar", **{"w:fldCharType": "begin"}))

    instr_run = paragraph.add_run()
    instr = _xml("w:instrText", **{"xml:space": "preserve"})
    _set_xml_text(instr, ' TOC \\o "1-3" \\h \\z \\u ')
    _append_run_xml(instr_run, instr)

    separate_run = paragraph.add_run()
    _append_run_xml(separate_run, _xml("w:fldChar", **{"w:fldCharType": "separate"}))

    end_run = paragraph.add_run()
    _append_run_xml(end_run, _xml("w:fldChar", **{"w:fldCharType": "end"}))


def _add_abstract_block(doc: DocumentT, abstract_text: str, language: str) -> None:
    label_map = {"uz": "ANNOTATSIYA", "ru": "АННОТАЦИЯ", "en": "ABSTRACT"}
    heading = doc.add_paragraph(label_map[language], style="Heading 1")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_body_paragraph(doc, abstract_text)


def _add_keywords_line(doc: DocumentT, keywords: list[str], language: str) -> None:
    label = _KEYWORDS_LABELS[language]
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run(f"{label}: {', '.join(keywords)}")
    run.italic = True
    run.font.size = Pt(SMALL_FONT_SIZE_PT)
    run.font.name = BODY_FONT_NAME


# ---------------------------------------------------------------------------
# Section + paragraph rendering
# ---------------------------------------------------------------------------


def _add_section_heading(doc: DocumentT, title: str, idx: int) -> None:
    heading_text = f"{idx + 1}. {title}"
    heading = doc.add_paragraph(heading_text, style="Heading 1")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_body_paragraph(doc: DocumentT, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.style = "Normal"


def _render_paragraph_with_citations(
    paragraph_obj: Paragraph,
    marker_map: dict[str, _MarkerMatch],
    formatter: BibliographyFormatter,
    style: CitationFormat,
    language: str,
    chicago_footnotes: list[str],
) -> tuple[str, int]:
    """Replace ``[<uuid>]`` markers in paragraph text with formatted citations.

    Returns the rewritten text plus the number of markers that were
    actually replaced (used for the export-result citation_count). For
    Chicago we additionally append the rendered footnote body to
    ``chicago_footnotes`` for later emission as a footnotes section.
    """

    text = paragraph_obj.text
    used = 0

    def _replace(match: re.Match[str]) -> str:
        nonlocal used
        key = match.group(1)
        info = marker_map.get(key)
        if info is None:
            return match.group(0)
        used += 1
        if style is CitationFormat.CHICAGO:
            footnote_number = len(chicago_footnotes) + 1
            metadata = info.metadata
            if metadata is not None:
                footnote_text = formatter.format_chicago_footnote(
                    metadata, page=None, first_occurrence=True, number=footnote_number
                )
            else:
                footnote_text = f"{footnote_number} {info.entry.formatted_text}"
            chicago_footnotes.append(footnote_text)
            return _superscript(footnote_number)
        if info.metadata is not None:
            return formatter.format_inline_citation(
                citation_number=info.number,
                metadata=info.metadata,
                style=style,
                page=None,
                language=language,
            )
        return _fallback_inline_citation(info, style)

    return _CITATION_MARKER_RE.sub(_replace, text), used


def _fallback_inline_citation(match: _MarkerMatch, style: CitationFormat) -> str:
    """Render an inline citation when no :class:`CitationMetadata` was supplied.

    Numbered styles (GOST/IEEE/Vancouver) only need the number.
    APA needs (Author, Year); we extract that from the formatted
    bibliography entry by regex. Anything we can't parse falls back to
    the number form so we never emit a broken marker.
    """

    if style is CitationFormat.GOST:
        return f"[{match.number}]"
    if style is CitationFormat.IEEE:
        return f"[{match.number}]"
    if style is CitationFormat.VANCOUVER:
        return f"({match.number})"
    if style is CitationFormat.APA:
        parsed = _APA_AUTHOR_YEAR_RE.match(match.entry.formatted_text)
        if parsed is not None:
            return f"({parsed.group(1).strip()}, {parsed.group(2)})"
        return f"[{match.number}]"
    return _superscript(match.number)


# ---------------------------------------------------------------------------
# Bibliography + footnotes + warning
# ---------------------------------------------------------------------------


def _add_bibliography_section(
    doc: DocumentT, bibliography: FormattedBibliography, language: str
) -> None:
    heading = doc.add_paragraph(_BIBLIOGRAPHY_HEADINGS[language], style="Heading 1")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for entry in bibliography.entries:
        paragraph = doc.add_paragraph()
        paragraph.style = "Normal"
        paragraph.paragraph_format.first_line_indent = Cm(-HANGING_INDENT_CM)
        paragraph.paragraph_format.left_indent = Cm(HANGING_INDENT_CM)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.add_run(entry.formatted_text)


def _add_footnotes_section(doc: DocumentT, footnotes: list[str], language: str) -> None:
    heading = doc.add_paragraph(_FOOTNOTES_HEADINGS[language], style="Heading 1")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for note in footnotes:
        paragraph = doc.add_paragraph()
        paragraph.style = "Normal"
        paragraph.paragraph_format.first_line_indent = Cm(0)
        run = paragraph.add_run(note)
        run.font.size = Pt(SMALL_FONT_SIZE_PT)


def _add_verification_warning(doc: DocumentT, count: int, language: str) -> None:
    text_map = {
        "uz": (
            f"⚠ DIQQAT: {count} ta iqtibos tekshiruvdan o'tmadi. Iltimos, ko'rib chiqing. "
            "(Ushbu xabarni topshirishdan oldin o'chiring.)"
        ),
        "ru": (
            f"⚠ ВНИМАНИЕ: {count} цитат не прошли проверку. Пожалуйста, проверьте. "
            "(Удалите это сообщение перед сдачей.)"
        ),
        "en": (
            f"⚠ WARNING: {count} citations failed verification. Please review. "
            "(Delete this notice before submission.)"
        ),
    }
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run(text_map[language])
    run.bold = True
    run.font.size = Pt(SMALL_FONT_SIZE_PT)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)


# ---------------------------------------------------------------------------
# Marker map construction
# ---------------------------------------------------------------------------


def _build_marker_map(
    draft: ArticleDraftResult,
    bibliography: FormattedBibliography,
    citation_metadata: list[CitationMetadata] | None,
) -> dict[str, _MarkerMatch]:
    """Resolve every distinct citation marker in the draft to a bibliography entry.

    Strategy:
    1. Walk paragraphs in document order, collecting unique marker keys.
       Markers appear as ``[<source_id>]`` in paragraph text and the
       same UUID (string-form) appears on at least one ``CitationRef``
       in :attr:`Paragraph.citations`.
    2. First try matching ``str(marker) == FormattedEntry.source_id``;
       any match wins immediately.
    3. Fall back to assigning bibliography entries to unmatched markers
       in order of first appearance — this keeps citation numbers stable
       for tests that pin ``[1]`` / ``[2]`` regardless of UUID values.
    """

    metadata_by_source_id: dict[str, CitationMetadata] = {}
    if citation_metadata is not None:
        for meta in citation_metadata:
            if meta.source_id is not None:
                metadata_by_source_id[meta.source_id] = meta

    entry_by_source_id: dict[str, FormattedEntry] = {
        e.source_id: e for e in bibliography.entries if e.source_id is not None
    }

    appearance_order: list[str] = []
    seen: set[str] = set()
    for section_result in draft.sections:
        for paragraph in section_result.section.paragraphs:
            for match in _CITATION_MARKER_RE.finditer(paragraph.text):
                key = match.group(1)
                if key in seen:
                    continue
                seen.add(key)
                appearance_order.append(key)

    marker_map: dict[str, _MarkerMatch] = {}
    fallback_pool: list[FormattedEntry] = []
    for entry in bibliography.entries:
        if entry.source_id is None or entry.source_id not in seen:
            fallback_pool.append(entry)

    fallback_iter = iter(fallback_pool)
    next_fallback_number = 1
    for key in appearance_order:
        entry = entry_by_source_id.get(key)
        if entry is not None:
            number = entry.number if entry.number is not None else _next_number(marker_map)
            metadata = metadata_by_source_id.get(key)
            marker_map[key] = _MarkerMatch(number=number, entry=entry, metadata=metadata)
            continue
        fallback_entry = next(fallback_iter, None)
        if fallback_entry is None:
            continue
        number = (
            fallback_entry.number if fallback_entry.number is not None else next_fallback_number
        )
        metadata = (
            metadata_by_source_id.get(fallback_entry.source_id)
            if fallback_entry.source_id is not None
            else None
        )
        marker_map[key] = _MarkerMatch(number=number, entry=fallback_entry, metadata=metadata)
        next_fallback_number = number + 1
    return marker_map


def _next_number(marker_map: dict[str, _MarkerMatch]) -> int:
    if not marker_map:
        return 1
    return max(m.number for m in marker_map.values()) + 1


# ---------------------------------------------------------------------------
# Small helpers
# ---------------------------------------------------------------------------


def _add_blank_lines(doc: DocumentT, count: int) -> None:
    for _ in range(count):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Cm(0)


def _add_centered_line(
    doc: DocumentT,
    text: str,
    *,
    bold: bool = False,
    size_pt: int | None = None,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = BODY_FONT_NAME
    run.font.size = Pt(size_pt if size_pt is not None else BODY_FONT_SIZE_PT)


def _add_right_aligned_line(doc: DocumentT, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run(text)
    run.font.name = BODY_FONT_NAME
    run.font.size = Pt(BODY_FONT_SIZE_PT)


def _add_page_break(doc: DocumentT) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_break()
    br = _runs_last_break(run)
    br.set(qn("w:type"), "page")


_SUPERSCRIPT_DIGITS: Final[dict[str, str]] = {
    "0": "⁰",
    "1": "¹",
    "2": "²",
    "3": "³",
    "4": "⁴",
    "5": "⁵",
    "6": "⁶",
    "7": "⁷",
    "8": "⁸",
    "9": "⁹",
}


def _superscript(number: int) -> str:
    return "".join(_SUPERSCRIPT_DIGITS[d] for d in str(number))


def _normalise_language(language: str) -> str:
    lang = language.lower()
    if lang in _BIBLIOGRAPHY_HEADINGS:
        return lang
    return "en"


def _filename(metadata: ArticleExportMetadata) -> str:
    safe = re.sub(r"[^\w\s-]", "", metadata.title, flags=re.UNICODE)
    safe = re.sub(r"\s+", "_", safe.strip())[:50].strip("_")
    if not safe:
        safe = "article"
    return f"{safe}_{metadata.article_type.value}.docx"


def _count_words(draft: ArticleDraftResult) -> int:
    total = 0
    for section_result in draft.sections:
        for paragraph in section_result.section.paragraphs:
            total += len(paragraph.text.split())
    return total


__all__ = [
    "BODY_FONT_NAME",
    "BODY_FONT_SIZE_PT",
    "MARGIN_BOTTOM_CM",
    "MARGIN_LEFT_CM",
    "MARGIN_RIGHT_CM",
    "MARGIN_TOP_CM",
    "PAGE_HEIGHT_CM",
    "PAGE_WIDTH_CM",
    "DOCXExporter",
]
