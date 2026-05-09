"""DOCX parser built on python-docx.

Word documents have no real pages, so we group consecutive non-heading
paragraphs into one logical "page" rooted at the preceding heading. If the
document has no headings, the whole body is a single page. All tables in the
document are attached to the first page that has any content.
"""

from __future__ import annotations

import asyncio
import logging
import re
from io import BytesIO
from typing import Final

from docx import Document
from docx.document import Document as DocxDocument

from packages.core.models.source import (
    ParsedPage,
    ParsedSource,
    SourceMetadataExtracted,
)
from packages.workers.source.parsers.lang_detect import detect_language

logger = logging.getLogger(__name__)


DOI_RE: Final[re.Pattern[str]] = re.compile(
    r"10\.\d{4,9}/[-._;()/:A-Z0-9]+",
    re.IGNORECASE,
)


class DOCXParser:
    """Parses a DOCX byte stream into a :class:`ParsedSource`."""

    async def parse(self, file_bytes: bytes, filename: str) -> ParsedSource:
        return await asyncio.to_thread(self._parse_sync, file_bytes, filename)

    def _parse_sync(self, file_bytes: bytes, filename: str) -> ParsedSource:
        errors: list[str] = []

        try:
            doc = Document(BytesIO(file_bytes))
        except Exception as exc:
            logger.warning("Failed to open DOCX %s: %s", filename, exc)
            return ParsedSource(
                filename=filename,
                file_type="docx",
                file_size_bytes=len(file_bytes),
                pages=[],
                metadata=SourceMetadataExtracted(),
                full_text="",
                needs_ocr_pages=[],
                parse_errors=[f"Failed to open DOCX: {exc}"],
            )

        pages = self._build_pages(doc, errors)
        if not pages:
            pages = [
                ParsedPage(
                    page_number=1,
                    text="",
                    char_count=0,
                    needs_ocr=False,
                )
            ]

        tables = self._extract_tables(doc, errors)
        if tables:
            pages[0] = pages[0].model_copy(update={"tables": tables})

        full_text = "\n\n".join(p.text for p in pages if p.text)
        word_count = len(full_text.split()) if full_text else 0
        metadata = self._build_metadata(doc, full_text, word_count, len(pages))

        return ParsedSource(
            filename=filename,
            file_type="docx",
            file_size_bytes=len(file_bytes),
            pages=pages,
            metadata=metadata,
            full_text=full_text,
            needs_ocr_pages=[],
            parse_errors=errors,
        )

    @staticmethod
    def _build_pages(doc: DocxDocument, errors: list[str]) -> list[ParsedPage]:
        pages: list[ParsedPage] = []
        current_heading: str | None = None
        current_buffer: list[str] = []
        current_headings: list[str] = []

        def flush() -> None:
            text = "\n".join(current_buffer).strip()
            if not text and not current_headings:
                return
            page_text_parts: list[str] = []
            if current_heading:
                page_text_parts.append(current_heading)
            if text:
                page_text_parts.append(text)
            full = "\n\n".join(page_text_parts)
            char_count = len(re.sub(r"\s+", "", full))
            pages.append(
                ParsedPage(
                    page_number=len(pages) + 1,
                    text=full,
                    char_count=char_count,
                    needs_ocr=False,
                    headings=list(current_headings),
                    tables=[],
                )
            )

        try:
            for paragraph in doc.paragraphs:
                style_name = (paragraph.style.name or "") if paragraph.style else ""
                text = paragraph.text or ""
                if style_name.startswith("Heading") or style_name in {"Title", "Subtitle"}:
                    flush()
                    current_buffer = []
                    current_heading = text.strip() or None
                    current_headings = [text.strip()] if text.strip() else []
                else:
                    if text.strip():
                        current_buffer.append(text)
            flush()
        except Exception as exc:
            errors.append(f"Paragraph walk failed: {exc}")
            logger.warning("DOCX paragraph walk failed: %s", exc)

        return pages

    @staticmethod
    def _extract_tables(doc: DocxDocument, errors: list[str]) -> list[list[list[str]]]:
        out: list[list[list[str]]] = []
        try:
            for table in doc.tables:
                rows: list[list[str]] = []
                for row in table.rows:
                    rows.append([cell.text.strip() for cell in row.cells])
                if rows:
                    out.append(rows)
        except Exception as exc:
            errors.append(f"Table extraction failed: {exc}")
            logger.warning("DOCX table extraction failed: %s", exc)
        return out

    @staticmethod
    def _build_metadata(
        doc: DocxDocument,
        full_text: str,
        word_count: int,
        page_count: int,
    ) -> SourceMetadataExtracted:
        props = doc.core_properties
        title = (props.title or "").strip() or None
        author_field = (props.author or "").strip()
        authors = [a.strip() for a in re.split(r"[,;]", author_field) if a.strip()]
        created = props.created
        creation_date = created.isoformat() if created else None
        year = created.year if created else None

        doi_match = DOI_RE.search(full_text[:500])
        doi = doi_match.group(0) if doi_match else None

        return SourceMetadataExtracted(
            title=title,
            authors=authors,
            year=year,
            doi=doi,
            page_count=page_count,
            word_count=word_count,
            language_detected=detect_language(full_text),
            has_images=False,
            creation_date=creation_date,
        )
