"""Tests for :class:`DOCXExporter`.

The exporter is wrapped python-docx code: tests run the real exporter
end-to-end, then re-open the produced bytes with python-docx to assert
on what Word will actually see (page setup, fonts, paragraph text,
bibliography heading, footer field codes, etc.). No python-docx mocks,
per ``.claude/rules/testing.md``.
"""

from __future__ import annotations

from datetime import UTC, datetime
from io import BytesIO
from uuid import UUID, uuid4

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from packages.core.enums import (
    ArticleSectionStatus,
    ArticleStructure,
    CitationFormat,
    CitationVerdict,
    SourceType,
)
from packages.core.models.article import (
    ArticleDraftResult,
    ArticleOutline,
    ArticleQualitySummary,
    ArticleSection,
    CitationRef,
    DraftResult,
    OutlineSection,
    Paragraph,
    QualityCheckResult,
)
from packages.core.models.bibliography import (
    CitationMetadata,
    FormattedBibliography,
)
from packages.core.models.export import ArticleExportMetadata, ExportResult
from packages.core.models.verification import (
    CitationVerification,
    CitationVerificationReport,
)
from packages.workers.article.docx_export import (
    BODY_FONT_NAME,
    BODY_FONT_SIZE_PT,
    DOCXExporter,
)

# ---------------------------------------------------------------------------
# Fixture helpers
# ---------------------------------------------------------------------------


def _make_section(
    article_id: UUID,
    index: int,
    title: str,
    paragraphs: list[Paragraph],
) -> ArticleSection:
    return ArticleSection(
        article_id=article_id,
        section_index=index,
        title=title,
        paragraphs=paragraphs,
        word_count=sum(len(p.text.split()) for p in paragraphs),
        status=ArticleSectionStatus.DRAFT,
        created_at=datetime.now(UTC),
    )


def _make_draft_result(section: ArticleSection) -> DraftResult:
    return DraftResult(
        section=section,
        quality_check=QualityCheckResult(
            passed=True, checks_passed=["ok"], checks_failed=[], overall_score=1.0
        ),
    )


def _make_draft(sections: list[ArticleSection]) -> ArticleDraftResult:
    summary = ArticleQualitySummary(
        sections_passed=len(sections),
        sections_failed=0,
        sections_revised=0,
        overall_score=1.0,
        weakest_section="",
        strongest_section="",
    )
    word_count = sum(s.word_count for s in sections)
    return ArticleDraftResult(
        sections=[_make_draft_result(s) for s in sections],
        total_word_count=word_count,
        total_llm_calls=len(sections),
        total_tokens=100 * len(sections),
        estimated_cost_usd=0.01,
        quality_summary=summary,
        warnings=[],
    )


def _make_outline(structure: ArticleStructure, section_titles: list[str]) -> ArticleOutline:
    return ArticleOutline(
        title="Test Article",
        structure=structure,
        sections=[
            OutlineSection(title=t, target_words=200, purpose="purpose") for t in section_titles
        ],
        thesis="A non-trivial thesis statement that meets the length requirement.",
        total_target_words=200 * len(section_titles),
    )


def _make_metadata(
    article_type: ArticleStructure = ArticleStructure.REFERAT,
    citation_format: CitationFormat = CitationFormat.GOST,
    *,
    title: str = "Test Article",
    university_name: str | None = "Toshkent davlat universiteti",
    supervisor_name: str | None = "Karimov A.",
    keywords: list[str] | None = None,
    abstract_text: str | None = None,
) -> ArticleExportMetadata:
    return ArticleExportMetadata(
        title=title,
        author_name="Tursunov B.",
        author_group="IF-22-01",
        supervisor_name=supervisor_name,
        supervisor_title="dots., professor",
        university_name=university_name,
        faculty_name="Informatika fakulteti",
        department_name="Dasturlash kafedrasi",
        city="Toshkent",
        year=2026,
        article_type=article_type,
        citation_format=citation_format,
        keywords=keywords or [],
        abstract_text=abstract_text,
    )


def _journal_meta(source_id: str, number: int = 1) -> CitationMetadata:
    return CitationMetadata(
        title="Passive radiative cooling",
        authors=["Raman, Aaswath P.", "Anoma, Marc Abou"],
        year=2014,
        journal="Nature",
        volume="515",
        pages="540-544",
        source_type=SourceType.JOURNAL_ARTICLE,
        source_id=source_id,
        citation_number=number,
    )


def _make_bibliography(
    entries_meta: list[CitationMetadata],
    style: CitationFormat = CitationFormat.GOST,
    language: str = "uz",
) -> FormattedBibliography:
    """Build a FormattedBibliography by running the real BibliographyFormatter."""

    from packages.workers.article.bibliography import BibliographyFormatter

    formatter = BibliographyFormatter()
    return formatter.format_bibliography(entries_meta, style=style, language=language)


def _open_docx(file_bytes: bytes) -> Document:  # type: ignore[valid-type]
    return Document(BytesIO(file_bytes))


def _all_paragraph_text(doc: Document) -> list[str]:  # type: ignore[valid-type]
    return [p.text for p in doc.paragraphs]


def _full_text(doc: Document) -> str:  # type: ignore[valid-type]
    return "\n".join(p.text for p in doc.paragraphs)


def _heading_paragraphs(doc: Document, level: int) -> list[str]:  # type: ignore[valid-type]
    target = f"Heading {level}"
    return [p.text for p in doc.paragraphs if p.style.name == target]


# ---------------------------------------------------------------------------
# Document setup tests
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_export_produces_valid_docx() -> None:
    article_id = uuid4()
    section = _make_section(
        article_id, 0, "Kirish", [Paragraph(text="Birinchi paragraf matni.", citations=[])]
    )
    draft = _make_draft([section])
    outline = _make_outline(ArticleStructure.REFERAT, ["Kirish"])
    bibliography = _make_bibliography([])
    metadata = _make_metadata()

    result = await DOCXExporter().export(draft, bibliography, None, outline, metadata, "uz")

    assert result.file_bytes
    doc = _open_docx(result.file_bytes)
    assert len(doc.paragraphs) > 0


@pytest.mark.asyncio
async def test_export_page_setup_a4() -> None:
    article_id = uuid4()
    draft = _make_draft(
        [_make_section(article_id, 0, "Kirish", [Paragraph(text="Matn", citations=[])])]
    )
    outline = _make_outline(ArticleStructure.REFERAT, ["Kirish"])
    metadata = _make_metadata()

    result = await DOCXExporter().export(
        draft, _make_bibliography([]), None, outline, metadata, "uz"
    )
    doc = _open_docx(result.file_bytes)
    section = doc.sections[0]
    # Page dimensions are stored as twips internally; allow ~1 twip rounding (~635 EMU)
    assert abs(section.page_width - Cm(21.0)) < 1000
    assert abs(section.page_height - Cm(29.7)) < 1000


@pytest.mark.asyncio
async def test_export_margins_correct() -> None:
    article_id = uuid4()
    draft = _make_draft(
        [_make_section(article_id, 0, "Kirish", [Paragraph(text="Matn", citations=[])])]
    )
    metadata = _make_metadata()
    outline = _make_outline(ArticleStructure.REFERAT, ["Kirish"])

    result = await DOCXExporter().export(
        draft, _make_bibliography([]), None, outline, metadata, "uz"
    )
    section = _open_docx(result.file_bytes).sections[0]
    # Margins are stored as twips internally; allow ~1 twip rounding tolerance
    assert abs(section.top_margin - Cm(2.0)) < 1000
    assert abs(section.bottom_margin - Cm(2.0)) < 1000
    assert abs(section.left_margin - Cm(3.0)) < 1000
    assert abs(section.right_margin - Cm(1.5)) < 1000


@pytest.mark.asyncio
async def test_export_default_font() -> None:
    article_id = uuid4()
    draft = _make_draft(
        [_make_section(article_id, 0, "Kirish", [Paragraph(text="Matn", citations=[])])]
    )
    outline = _make_outline(ArticleStructure.REFERAT, ["Kirish"])
    metadata = _make_metadata()
    result = await DOCXExporter().export(
        draft, _make_bibliography([]), None, outline, metadata, "uz"
    )
    doc = _open_docx(result.file_bytes)
    style = doc.styles["Normal"]
    assert style.font.name == BODY_FONT_NAME
    assert style.font.size == Pt(BODY_FONT_SIZE_PT)


@pytest.mark.asyncio
async def test_export_line_spacing() -> None:
    article_id = uuid4()
    draft = _make_draft(
        [_make_section(article_id, 0, "Kirish", [Paragraph(text="Matn", citations=[])])]
    )
    outline = _make_outline(ArticleStructure.REFERAT, ["Kirish"])
    metadata = _make_metadata()
    result = await DOCXExporter().export(
        draft, _make_bibliography([]), None, outline, metadata, "uz"
    )
    doc = _open_docx(result.file_bytes)
    spacing = doc.styles["Normal"].paragraph_format.line_spacing_rule
    assert spacing == WD_LINE_SPACING.ONE_POINT_FIVE


@pytest.mark.asyncio
async def test_export_first_line_indent() -> None:
    article_id = uuid4()
    draft = _make_draft(
        [_make_section(article_id, 0, "Kirish", [Paragraph(text="Matn", citations=[])])]
    )
    outline = _make_outline(ArticleStructure.REFERAT, ["Kirish"])
    metadata = _make_metadata()
    result = await DOCXExporter().export(
        draft, _make_bibliography([]), None, outline, metadata, "uz"
    )
    doc = _open_docx(result.file_bytes)
    indent = doc.styles["Normal"].paragraph_format.first_line_indent
    # twips rounding tolerance, same as page dimensions above
    assert abs(indent - Cm(1.25)) < 1000


@pytest.mark.asyncio
async def test_export_justified_alignment() -> None:
    article_id = uuid4()
    draft = _make_draft(
        [_make_section(article_id, 0, "Kirish", [Paragraph(text="Matn", citations=[])])]
    )
    outline = _make_outline(ArticleStructure.REFERAT, ["Kirish"])
    metadata = _make_metadata()
    result = await DOCXExporter().export(
        draft, _make_bibliography([]), None, outline, metadata, "uz"
    )
    doc = _open_docx(result.file_bytes)
    alignment = doc.styles["Normal"].paragraph_format.alignment
    assert alignment == WD_ALIGN_PARAGRAPH.JUSTIFY


# ---------------------------------------------------------------------------
# Title page tests
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_export_referat_has_title_page() -> None:
    article_id = uuid4()
    draft = _make_draft(
        [_make_section(article_id, 0, "Kirish", [Paragraph(text="Matn", citations=[])])]
    )
    outline = _make_outline(ArticleStructure.REFERAT, ["Kirish"])
    metadata = _make_metadata(article_type=ArticleStructure.REFERAT)
    result = await DOCXExporter().export(
        draft, _make_bibliography([]), None, outline, metadata, "uz"
    )
    full_text = _full_text(_open_docx(result.file_bytes))
    assert "TOSHKENT DAVLAT UNIVERSITETI" in full_text
    assert "REFERAT" in full_text
    assert "Tursunov B." in full_text
    assert "Karimov A." in full_text


@pytest.mark.asyncio
async def test_export_kurs_ishi_has_title_page() -> None:
    article_id = uuid4()
    draft = _make_draft(
        [_make_section(article_id, 0, "Kirish", [Paragraph(text="Matn", citations=[])])]
    )
    outline = _make_outline(ArticleStructure.KURS_ISHI, ["Kirish"])
    metadata = _make_metadata(article_type=ArticleStructure.KURS_ISHI)
    result = await DOCXExporter().export(
        draft, _make_bibliography([]), None, outline, metadata, "uz"
    )
    full_text = _full_text(_open_docx(result.file_bytes))
    assert "KURS ISHI" in full_text


@pytest.mark.asyncio
async def test_export_ilmiy_maqola_no_title_page() -> None:
    article_id = uuid4()
    draft = _make_draft(
        [
            _make_section(
                article_id, 0, "Kirish", [Paragraph(text="Boshlanish matni.", citations=[])]
            )
        ]
    )
    outline = _make_outline(ArticleStructure.ILMIY_MAQOLA, ["Kirish"])
    metadata = _make_metadata(article_type=ArticleStructure.ILMIY_MAQOLA)
    result = await DOCXExporter().export(
        draft, _make_bibliography([]), None, outline, metadata, "uz"
    )
    full_text = _full_text(_open_docx(result.file_bytes))
    assert "ILMIY MAQOLA" not in full_text  # no type label
    # First non-empty paragraph should not be the university header
    first_real = next((line for line in full_text.split("\n") if line.strip()), "")
    assert "UNIVERSITETI" not in first_real


@pytest.mark.asyncio
async def test_export_title_page_without_optional_metadata() -> None:
    article_id = uuid4()
    draft = _make_draft(
        [_make_section(article_id, 0, "Kirish", [Paragraph(text="Matn", citations=[])])]
    )
    outline = _make_outline(ArticleStructure.REFERAT, ["Kirish"])
    metadata = _make_metadata(university_name=None, supervisor_name=None)
    result = await DOCXExporter().export(
        draft, _make_bibliography([]), None, outline, metadata, "uz"
    )
    full_text = _full_text(_open_docx(result.file_bytes))
    # Still has type label and author
    assert "REFERAT" in full_text
    assert "Tursunov B." in full_text
    # No supervisor block
    assert "Karimov" not in full_text


# ---------------------------------------------------------------------------
# Headings tests
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_export_heading_1_formatting() -> None:
    article_id = uuid4()
    draft = _make_draft(
        [_make_section(article_id, 0, "Kirish", [Paragraph(text="Matn", citations=[])])]
    )
    outline = _make_outline(ArticleStructure.REFERAT, ["Kirish"])
    metadata = _make_metadata()
    result = await DOCXExporter().export(
        draft, _make_bibliography([]), None, outline, metadata, "uz"
    )
    doc = _open_docx(result.file_bytes)
    style = doc.styles["Heading 1"]
    assert style.font.name == BODY_FONT_NAME
    assert style.font.bold is True
    assert style.font.size == Pt(BODY_FONT_SIZE_PT)


@pytest.mark.asyncio
async def test_export_heading_numbering() -> None:
    article_id = uuid4()
    sections = [
        _make_section(
            article_id, i, title, [Paragraph(text=f"Section {i + 1} matni.", citations=[])]
        )
        for i, title in enumerate(["Kirish", "Asosiy qism", "Xulosa"])
    ]
    draft = _make_draft(sections)
    outline = _make_outline(ArticleStructure.REFERAT, [s.title for s in sections])
    metadata = _make_metadata()
    result = await DOCXExporter().export(
        draft, _make_bibliography([]), None, outline, metadata, "uz"
    )
    doc = _open_docx(result.file_bytes)
    headings = _heading_paragraphs(doc, 1)
    assert "1. Kirish" in headings
    assert "2. Asosiy qism" in headings
    assert "3. Xulosa" in headings


# ---------------------------------------------------------------------------
# Content tests
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_export_paragraph_text_present() -> None:
    article_id = uuid4()
    sentence = "Bu juda muhim ilmiy paragraf matni."
    draft = _make_draft(
        [_make_section(article_id, 0, "Kirish", [Paragraph(text=sentence, citations=[])])]
    )
    outline = _make_outline(ArticleStructure.REFERAT, ["Kirish"])
    metadata = _make_metadata()
    result = await DOCXExporter().export(
        draft, _make_bibliography([]), None, outline, metadata, "uz"
    )
    assert sentence in _full_text(_open_docx(result.file_bytes))


@pytest.mark.asyncio
async def test_export_multiple_sections_in_order() -> None:
    article_id = uuid4()
    titles = ["Kirish", "Adabiyotlar tahlili", "Metodologiya", "Natijalar", "Xulosa"]
    sections = [
        _make_section(article_id, i, t, [Paragraph(text=f"{t} matni.", citations=[])])
        for i, t in enumerate(titles)
    ]
    draft = _make_draft(sections)
    outline = _make_outline(ArticleStructure.REFERAT, titles)
    metadata = _make_metadata()
    result = await DOCXExporter().export(
        draft, _make_bibliography([]), None, outline, metadata, "uz"
    )
    doc = _open_docx(result.file_bytes)
    headings = _heading_paragraphs(doc, 1)
    main_section_headings = [h for h in headings if any(t in h for t in titles)]
    assert main_section_headings == [f"{i + 1}. {t}" for i, t in enumerate(titles)]


# ---------------------------------------------------------------------------
# Citations tests
# ---------------------------------------------------------------------------


def _section_with_citation(
    article_id: UUID, marker: str, claim_id: UUID, source_id: UUID
) -> ArticleSection:
    text = f"This finding is well documented [{marker}]."
    return _make_section(
        article_id,
        0,
        "Findings",
        [Paragraph(text=text, citations=[CitationRef(claim_id=claim_id, source_id=source_id)])],
    )


@pytest.mark.asyncio
async def test_export_gost_inline_citations() -> None:
    article_id = uuid4()
    marker = "src_a"
    section = _section_with_citation(article_id, marker, uuid4(), uuid4())
    draft = _make_draft([section])
    outline = _make_outline(ArticleStructure.REFERAT, ["Findings"])
    bibliography = _make_bibliography([_journal_meta(marker, number=1)], style=CitationFormat.GOST)
    metadata = _make_metadata(citation_format=CitationFormat.GOST)
    result = await DOCXExporter().export(draft, bibliography, None, outline, metadata, "uz")
    full_text = _full_text(_open_docx(result.file_bytes))
    assert "[1]" in full_text
    assert f"[{marker}]" not in full_text


@pytest.mark.asyncio
async def test_export_apa_inline_citations() -> None:
    article_id = uuid4()
    marker = "src_apa"
    section = _section_with_citation(article_id, marker, uuid4(), uuid4())
    draft = _make_draft([section])
    outline = _make_outline(ArticleStructure.REFERAT, ["Findings"])
    bibliography = _make_bibliography(
        [_journal_meta(marker, number=1)], style=CitationFormat.APA, language="en"
    )
    metadata = _make_metadata(citation_format=CitationFormat.APA)
    result = await DOCXExporter().export(draft, bibliography, None, outline, metadata, "en")
    full_text = _full_text(_open_docx(result.file_bytes))
    assert "(Raman" in full_text
    assert "2014" in full_text


@pytest.mark.asyncio
async def test_export_ieee_inline_citations() -> None:
    article_id = uuid4()
    marker = "src_ieee"
    section = _section_with_citation(article_id, marker, uuid4(), uuid4())
    draft = _make_draft([section])
    outline = _make_outline(ArticleStructure.REFERAT, ["Findings"])
    bibliography = _make_bibliography(
        [_journal_meta(marker, number=1)], style=CitationFormat.IEEE, language="en"
    )
    metadata = _make_metadata(citation_format=CitationFormat.IEEE)
    result = await DOCXExporter().export(draft, bibliography, None, outline, metadata, "en")
    full_text = _full_text(_open_docx(result.file_bytes))
    assert "[1]" in full_text


# ---------------------------------------------------------------------------
# Bibliography tests
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_export_bibliography_section_present() -> None:
    article_id = uuid4()
    draft = _make_draft(
        [_make_section(article_id, 0, "Kirish", [Paragraph(text="Matn", citations=[])])]
    )
    outline = _make_outline(ArticleStructure.REFERAT, ["Kirish"])
    bibliography = _make_bibliography(
        [_journal_meta("a", 1), _journal_meta("b", 2)], style=CitationFormat.GOST
    )
    metadata = _make_metadata()
    result = await DOCXExporter().export(draft, bibliography, None, outline, metadata, "uz")
    full_text = _full_text(_open_docx(result.file_bytes))
    assert "FOYDALANILGAN ADABIYOTLAR RO'YXATI" in full_text
    bib_pos = full_text.find("FOYDALANILGAN")
    section_pos = full_text.find("1. Kirish")
    assert bib_pos > section_pos


@pytest.mark.asyncio
async def test_export_bibliography_heading_uzbek() -> None:
    article_id = uuid4()
    draft = _make_draft(
        [_make_section(article_id, 0, "Kirish", [Paragraph(text="Matn", citations=[])])]
    )
    outline = _make_outline(ArticleStructure.REFERAT, ["Kirish"])
    metadata = _make_metadata()
    result = await DOCXExporter().export(
        draft, _make_bibliography([]), None, outline, metadata, "uz"
    )
    full_text = _full_text(_open_docx(result.file_bytes))
    assert "FOYDALANILGAN ADABIYOTLAR RO'YXATI" in full_text


@pytest.mark.asyncio
async def test_export_bibliography_heading_russian() -> None:
    article_id = uuid4()
    draft = _make_draft(
        [_make_section(article_id, 0, "Введение", [Paragraph(text="Текст", citations=[])])]
    )
    outline = _make_outline(ArticleStructure.REFERAT, ["Введение"])
    metadata = _make_metadata()
    result = await DOCXExporter().export(
        draft, _make_bibliography([]), None, outline, metadata, "ru"
    )
    full_text = _full_text(_open_docx(result.file_bytes))
    assert "СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ" in full_text


@pytest.mark.asyncio
async def test_export_bibliography_heading_english() -> None:
    article_id = uuid4()
    draft = _make_draft(
        [_make_section(article_id, 0, "Introduction", [Paragraph(text="Text", citations=[])])]
    )
    outline = _make_outline(ArticleStructure.REFERAT, ["Introduction"])
    metadata = _make_metadata()
    result = await DOCXExporter().export(
        draft, _make_bibliography([]), None, outline, metadata, "en"
    )
    full_text = _full_text(_open_docx(result.file_bytes))
    assert "REFERENCES" in full_text


@pytest.mark.asyncio
async def test_export_bibliography_entries_formatted() -> None:
    article_id = uuid4()
    draft = _make_draft(
        [_make_section(article_id, 0, "Kirish", [Paragraph(text="Matn", citations=[])])]
    )
    outline = _make_outline(ArticleStructure.REFERAT, ["Kirish"])
    metas = [_journal_meta(f"src_{i}", number=i + 1) for i in range(5)]
    bibliography = _make_bibliography(metas, style=CitationFormat.GOST)
    metadata = _make_metadata()
    result = await DOCXExporter().export(draft, bibliography, None, outline, metadata, "uz")
    full_text = _full_text(_open_docx(result.file_bytes))
    assert sum(1 for entry in bibliography.entries if "Raman" in entry.formatted_text) == 5
    assert full_text.count("Raman") == 5


# ---------------------------------------------------------------------------
# Keywords tests
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_export_keywords_for_ilmiy_maqola() -> None:
    article_id = uuid4()
    draft = _make_draft(
        [_make_section(article_id, 0, "Kirish", [Paragraph(text="Matn", citations=[])])]
    )
    outline = _make_outline(ArticleStructure.ILMIY_MAQOLA, ["Kirish"])
    metadata = _make_metadata(
        article_type=ArticleStructure.ILMIY_MAQOLA,
        keywords=["sovutish", "radiyatsiya", "samaradorlik"],
    )
    result = await DOCXExporter().export(
        draft, _make_bibliography([]), None, outline, metadata, "uz"
    )
    full_text = _full_text(_open_docx(result.file_bytes))
    assert "Kalit so'zlar" in full_text
    assert "sovutish" in full_text


@pytest.mark.asyncio
async def test_export_no_keywords_for_referat() -> None:
    article_id = uuid4()
    draft = _make_draft(
        [_make_section(article_id, 0, "Kirish", [Paragraph(text="Matn", citations=[])])]
    )
    outline = _make_outline(ArticleStructure.REFERAT, ["Kirish"])
    metadata = _make_metadata(
        article_type=ArticleStructure.REFERAT, keywords=["this", "should", "not", "appear"]
    )
    result = await DOCXExporter().export(
        draft, _make_bibliography([]), None, outline, metadata, "uz"
    )
    full_text = _full_text(_open_docx(result.file_bytes))
    assert "Kalit so'zlar" not in full_text


# ---------------------------------------------------------------------------
# Page numbers tests
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_export_has_page_numbers() -> None:
    article_id = uuid4()
    sections = [
        _make_section(article_id, i, t, [Paragraph(text=f"{t} matni.", citations=[])])
        for i, t in enumerate(["Kirish", "Asosiy", "Xulosa"])
    ]
    draft = _make_draft(sections)
    outline = _make_outline(ArticleStructure.REFERAT, ["Kirish", "Asosiy", "Xulosa"])
    metadata = _make_metadata()
    result = await DOCXExporter().export(
        draft, _make_bibliography([]), None, outline, metadata, "uz"
    )
    doc = _open_docx(result.file_bytes)
    footer = doc.sections[0].footer
    instr_texts = [
        t.text or ""
        for paragraph in footer.paragraphs
        for t in paragraph._element.iter(qn("w:instrText"))
    ]
    assert any("PAGE" in s for s in instr_texts)


# ---------------------------------------------------------------------------
# Verification warnings tests
# ---------------------------------------------------------------------------


def _make_verification(critical_count: int) -> CitationVerificationReport:
    verifications = [
        CitationVerification(
            section_id="s1",
            paragraph_index=0,
            citation_index=i,
            claim_id=f"claim_{i}",
            source_chunk_id=f"src_{i}",
            verdict=CitationVerdict.NOT_SUPPORTED,
            confidence=0.95,
            explanation="bad",
        )
        for i in range(critical_count)
    ]
    return CitationVerificationReport(
        total_citations=critical_count,
        supported=0,
        partially_supported=0,
        overclaimed=0,
        not_supported=critical_count,
        contradicted=0,
        source_not_found=0,
        overall_integrity_score=0.0,
        verifications=verifications,
        critical_issues=verifications,
        warnings=[],
    )


@pytest.mark.asyncio
async def test_export_with_clean_verification_no_warning() -> None:
    article_id = uuid4()
    draft = _make_draft(
        [_make_section(article_id, 0, "Kirish", [Paragraph(text="Matn", citations=[])])]
    )
    outline = _make_outline(ArticleStructure.REFERAT, ["Kirish"])
    metadata = _make_metadata()
    clean = CitationVerificationReport(
        total_citations=2,
        supported=2,
        partially_supported=0,
        overclaimed=0,
        not_supported=0,
        contradicted=0,
        source_not_found=0,
        overall_integrity_score=1.0,
    )
    result = await DOCXExporter().export(
        draft, _make_bibliography([]), clean, outline, metadata, "uz"
    )
    full_text = _full_text(_open_docx(result.file_bytes))
    assert "DIQQAT" not in full_text
    assert "WARNING" not in full_text


@pytest.mark.asyncio
async def test_export_with_critical_issues_shows_warning() -> None:
    article_id = uuid4()
    draft = _make_draft(
        [_make_section(article_id, 0, "Kirish", [Paragraph(text="Matn", citations=[])])]
    )
    outline = _make_outline(ArticleStructure.REFERAT, ["Kirish"])
    metadata = _make_metadata()
    report = _make_verification(critical_count=2)
    result = await DOCXExporter().export(
        draft, _make_bibliography([]), report, outline, metadata, "uz"
    )
    doc = _open_docx(result.file_bytes)
    full_text = _full_text(doc)
    assert "DIQQAT" in full_text
    assert "2" in full_text
    # Verify red colour on the warning run
    warning_runs = [
        run for paragraph in doc.paragraphs for run in paragraph.runs if "DIQQAT" in run.text
    ]
    assert warning_runs
    rgb = warning_runs[0].font.color.rgb
    assert rgb is not None
    # red component (high) higher than blue (low)
    assert rgb[0] > rgb[2]


# ---------------------------------------------------------------------------
# File metadata tests
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_export_filename_sanitized() -> None:
    article_id = uuid4()
    draft = _make_draft(
        [_make_section(article_id, 0, "Kirish", [Paragraph(text="Matn", citations=[])])]
    )
    outline = _make_outline(ArticleStructure.REFERAT, ["Kirish"])
    metadata = _make_metadata(title="Test: Article (Draft) #1")
    result = await DOCXExporter().export(
        draft, _make_bibliography([]), None, outline, metadata, "uz"
    )
    assert result.filename.endswith(".docx")
    assert ":" not in result.filename
    assert "(" not in result.filename
    assert "#" not in result.filename


@pytest.mark.asyncio
async def test_export_word_count_correct() -> None:
    article_id = uuid4()
    paragraphs = [
        Paragraph(
            text="Bu yetti so'zli birinchi paragraf matni hisoblanadi.", citations=[]
        ),  # 7 words
        Paragraph(text="Ikkinchi besh so'zli paragraf.", citations=[]),  # 4 words
    ]
    section = _make_section(article_id, 0, "Kirish", paragraphs)
    draft = _make_draft([section])
    outline = _make_outline(ArticleStructure.REFERAT, ["Kirish"])
    metadata = _make_metadata()
    result = await DOCXExporter().export(
        draft, _make_bibliography([]), None, outline, metadata, "uz"
    )
    expected = sum(len(p.text.split()) for p in paragraphs)
    assert result.word_count == expected


@pytest.mark.asyncio
async def test_export_page_estimate_reasonable() -> None:
    article_id = uuid4()
    # Paragraph max_length is 5000 chars; split 3000 words across multiple paragraphs.
    paragraphs = [Paragraph(text=" ".join(["word"] * 600), citations=[]) for _ in range(5)]
    section = _make_section(article_id, 0, "Long", paragraphs)
    draft = _make_draft([section])
    outline = _make_outline(ArticleStructure.REFERAT, ["Long"])
    metadata = _make_metadata()
    result = await DOCXExporter().export(
        draft, _make_bibliography([]), None, outline, metadata, "uz"
    )
    assert 10 <= result.page_count_estimate <= 15


# ---------------------------------------------------------------------------
# Edge cases
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_export_empty_article() -> None:
    summary = ArticleQualitySummary(
        sections_passed=0,
        sections_failed=0,
        sections_revised=0,
        overall_score=0.0,
        weakest_section="",
        strongest_section="",
    )
    draft = ArticleDraftResult(
        sections=[],
        total_word_count=0,
        total_llm_calls=0,
        total_tokens=0,
        estimated_cost_usd=0.0,
        quality_summary=summary,
        warnings=[],
    )
    outline = ArticleOutline(
        title="Empty",
        structure=ArticleStructure.REFERAT,
        sections=[OutlineSection(title="Kirish", target_words=200, purpose="purpose")],
        thesis="A thesis with sufficient length to satisfy validation rules.",
        total_target_words=200,
    )
    metadata = _make_metadata()
    result = await DOCXExporter().export(
        draft, _make_bibliography([]), None, outline, metadata, "uz"
    )
    assert result.file_size_bytes > 0
    doc = _open_docx(result.file_bytes)
    assert len(doc.paragraphs) > 0


@pytest.mark.asyncio
async def test_export_very_long_article() -> None:
    article_id = uuid4()
    sections: list[ArticleSection] = []
    for i in range(10):
        paragraphs = [
            Paragraph(text=f"Paragraf {j + 1} bo'limining matni {i + 1}.", citations=[])
            for j in range(5)
        ]
        sections.append(_make_section(article_id, i, f"Bo'lim {i + 1}", paragraphs))
    draft = _make_draft(sections)
    outline = _make_outline(ArticleStructure.REFERAT, [s.title for s in sections])
    metadata = _make_metadata()
    result = await DOCXExporter().export(
        draft, _make_bibliography([]), None, outline, metadata, "uz"
    )
    assert result.section_count == 10
    doc = _open_docx(result.file_bytes)
    assert len(_heading_paragraphs(doc, 1)) >= 10  # at least 10 section headings


@pytest.mark.asyncio
async def test_export_cyrillic_text() -> None:
    article_id = uuid4()
    cyrillic_text = "Это очень важный научный параграф на русском языке."
    section = _make_section(
        article_id, 0, "Введение", [Paragraph(text=cyrillic_text, citations=[])]
    )
    draft = _make_draft([section])
    outline = _make_outline(ArticleStructure.REFERAT, ["Введение"])
    metadata = _make_metadata()
    result = await DOCXExporter().export(
        draft, _make_bibliography([]), None, outline, metadata, "ru"
    )
    assert cyrillic_text in _full_text(_open_docx(result.file_bytes))


@pytest.mark.asyncio
async def test_export_uzbek_special_chars() -> None:
    article_id = uuid4()
    uz_text = "O'zbek tilining lotin yozuvi va g'arbiy harflarni ishlatadi."
    section = _make_section(article_id, 0, "Kirish", [Paragraph(text=uz_text, citations=[])])
    draft = _make_draft([section])
    outline = _make_outline(ArticleStructure.REFERAT, ["Kirish"])
    metadata = _make_metadata()
    result = await DOCXExporter().export(
        draft, _make_bibliography([]), None, outline, metadata, "uz"
    )
    full_text = _full_text(_open_docx(result.file_bytes))
    assert "o'zbek" in full_text.lower()
    assert "g'arbiy" in full_text.lower()


@pytest.mark.asyncio
async def test_export_mixed_language() -> None:
    article_id = uuid4()
    section = _make_section(
        article_id,
        0,
        "Kirish",
        [
            Paragraph(
                text="Bu uzbek tilidagi paragraf, lekin manba inglizchada [src_a].", citations=[]
            )
        ],
    )
    draft = _make_draft([section])
    outline = _make_outline(ArticleStructure.REFERAT, ["Kirish"])
    bibliography = _make_bibliography(
        [_journal_meta("src_a", number=1)], style=CitationFormat.APA, language="en"
    )
    metadata = _make_metadata(citation_format=CitationFormat.APA)
    result = await DOCXExporter().export(draft, bibliography, None, outline, metadata, "uz")
    full_text = _full_text(_open_docx(result.file_bytes))
    assert "uzbek tilidagi" in full_text
    # APA inline citation should still render (English-author-year form)
    assert "Raman" in full_text


# ---------------------------------------------------------------------------
# Table of contents tests
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_export_kurs_ishi_has_toc() -> None:
    article_id = uuid4()
    draft = _make_draft(
        [_make_section(article_id, 0, "Kirish", [Paragraph(text="Matn", citations=[])])]
    )
    outline = _make_outline(ArticleStructure.KURS_ISHI, ["Kirish"])
    metadata = _make_metadata(article_type=ArticleStructure.KURS_ISHI)
    result = await DOCXExporter().export(
        draft, _make_bibliography([]), None, outline, metadata, "uz"
    )
    doc = _open_docx(result.file_bytes)
    instr_texts = [
        t.text or ""
        for paragraph in doc.paragraphs
        for t in paragraph._element.iter(qn("w:instrText"))
    ]
    assert any("TOC" in s for s in instr_texts)


@pytest.mark.asyncio
async def test_export_referat_no_toc() -> None:
    article_id = uuid4()
    draft = _make_draft(
        [_make_section(article_id, 0, "Kirish", [Paragraph(text="Matn", citations=[])])]
    )
    outline = _make_outline(ArticleStructure.REFERAT, ["Kirish"])
    metadata = _make_metadata(article_type=ArticleStructure.REFERAT)
    result = await DOCXExporter().export(
        draft, _make_bibliography([]), None, outline, metadata, "uz"
    )
    doc = _open_docx(result.file_bytes)
    instr_texts = [
        t.text or ""
        for paragraph in doc.paragraphs
        for t in paragraph._element.iter(qn("w:instrText"))
    ]
    assert not any("TOC" in s for s in instr_texts)


# ---------------------------------------------------------------------------
# Round-trip tests
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_export_roundtrip_read_all_text() -> None:
    article_id = uuid4()
    paragraphs = [
        Paragraph(text="Birinchi paragraf.", citations=[]),
        Paragraph(text="Ikkinchi paragraf.", citations=[]),
        Paragraph(text="Uchinchi paragraf.", citations=[]),
    ]
    section = _make_section(article_id, 0, "Kirish", paragraphs)
    draft = _make_draft([section])
    outline = _make_outline(ArticleStructure.REFERAT, ["Kirish"])
    metadata = _make_metadata()
    result = await DOCXExporter().export(
        draft, _make_bibliography([]), None, outline, metadata, "uz"
    )
    full_text = _full_text(_open_docx(result.file_bytes))
    for p in paragraphs:
        assert p.text in full_text


# ---------------------------------------------------------------------------
# Model validation tests
# ---------------------------------------------------------------------------


def test_article_export_metadata_model() -> None:
    metadata = ArticleExportMetadata(
        title="Test",
        author_name="Author",
        article_type=ArticleStructure.REFERAT,
        citation_format=CitationFormat.GOST,
    )
    dumped = metadata.model_dump()
    restored = ArticleExportMetadata.model_validate(dumped)
    assert restored == metadata


def test_export_result_model() -> None:
    result = ExportResult(
        file_bytes=b"PK\x03\x04 fake docx bytes",
        filename="test.docx",
        file_size_bytes=100,
        page_count_estimate=2,
        word_count=500,
        section_count=3,
        citation_count=5,
        bibliography_count=4,
        warnings=[],
    )
    assert result.file_size_bytes > 0
    dumped = result.model_dump()
    restored = ExportResult.model_validate(dumped)
    assert restored == result


def test_export_result_extra_forbid() -> None:
    from pydantic import ValidationError

    with pytest.raises(ValidationError):
        ExportResult.model_validate(
            {
                "file_bytes": b"x",
                "filename": "f.docx",
                "file_size_bytes": 1,
                "page_count_estimate": 1,
                "word_count": 1,
                "section_count": 1,
                "citation_count": 0,
                "bibliography_count": 0,
                "warnings": [],
                "unknown_extra_field": "x",
            }
        )
